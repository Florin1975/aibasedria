import os
import io
import re
import json
import logging
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Union, Optional, Tuple, BinaryIO
from enum import Enum
from dataclasses import dataclass, field
import tempfile
from pathlib import Path

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import optional dependencies with fallbacks
try:
    import PyPDF2
except ImportError:
    logger.warning("PyPDF2 not installed. PDF processing will be limited.")
    PyPDF2 = None

try:
    from pdf2image import convert_from_path, convert_from_bytes
except ImportError:
    logger.warning("pdf2image not installed. PDF to image conversion will be unavailable.")
    convert_from_path = None
    convert_from_bytes = None

try:
    import docx
except ImportError:
    logger.warning("python-docx not installed. DOCX processing will be unavailable.")
    docx = None

try:
    import pytesseract
except ImportError:
    logger.warning("pytesseract not installed. Tesseract OCR will be unavailable.")
    pytesseract = None

try:
    import easyocr
except ImportError:
    logger.warning("easyocr not installed. EasyOCR will be unavailable.")
    easyocr = None

try:
    import spacy
except ImportError:
    logger.warning("spacy not installed. NLP processing will be limited.")
    spacy = None

try:
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.cluster import KMeans
    from sklearn.naive_bayes import MultinomialNB
    from sklearn.svm import SVC
    from sklearn.pipeline import Pipeline
except ImportError:
    logger.warning("scikit-learn not installed. ML-based classification will be unavailable.")
    np = None
    TfidfVectorizer = None
    KMeans = None
    MultinomialNB = None
    SVC = None
    Pipeline = None

try:
    from transformers import pipeline
except ImportError:
    logger.warning("transformers not installed. Advanced AI features will be unavailable.")
    pipeline = None

try:
    from PIL import Image
except ImportError:
    logger.warning("PIL not installed. Image processing will be limited.")
    Image = None


class DocumentType(Enum):
    """Enum for document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Class for document metadata"""
    filename: str
    file_type: DocumentType
    page_count: int = 0
    author: str = ""
    creation_date: str = ""
    modified_date: str = ""
    title: str = ""
    size_bytes: int = 0
    custom_metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        result = {
            "filename": self.filename,
            "file_type": self.file_type.value,
            "page_count": self.page_count,
            "author": self.author,
            "creation_date": self.creation_date,
            "modified_date": self.modified_date,
            "title": self.title,
            "size_bytes": self.size_bytes,
        }
        if self.custom_metadata:
            result["custom_metadata"] = self.custom_metadata
        return result


@dataclass
class ProcessingResult:
    """Class for document processing results"""
    metadata: DocumentMetadata
    text_content: str
    pages: List[str] = field(default_factory=list)
    entities: Dict[str, List[str]] = field(default_factory=dict)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    form_fields: Dict[str, str] = field(default_factory=dict)
    classification: Dict[str, float] = field(default_factory=dict)
    summary: str = ""
    keywords: List[str] = field(default_factory=list)
    sentiment: Dict[str, float] = field(default_factory=dict)
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert processing result to dictionary"""
        result = {
            "metadata": self.metadata.to_dict(),
            "text_content": self.text_content,
        }
        if self.pages:
            result["pages"] = self.pages
        if self.entities:
            result["entities"] = self.entities
        if self.tables:
            result["tables"] = self.tables
        if self.form_fields:
            result["form_fields"] = self.form_fields
        if self.classification:
            result["classification"] = self.classification
        if self.summary:
            result["summary"] = self.summary
        if self.keywords:
            result["keywords"] = self.keywords
        if self.sentiment:
            result["sentiment"] = self.sentiment
        if self.error:
            result["error"] = self.error
        return result

    def to_json(self) -> str:
        """Convert processing result to JSON string"""
        return json.dumps(self.to_dict(), default=str)


class DocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    @abstractmethod
    def can_load(self, file_path: str) -> bool:
        """Check if the loader can handle this file type"""
        pass
    
    @abstractmethod
    def load(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Load document and return text content and metadata"""
        pass
    
    @abstractmethod
    def load_from_bytes(self, file_bytes: BinaryIO, filename: str) -> Tuple[str, DocumentMetadata]:
        """Load document from bytes and return text content and metadata"""
        pass


class PDFLoader(DocumentLoader):
    """PDF document loader"""
    
    def can_load(self, file_path: str) -> bool:
        """Check if the file is a PDF"""
        return file_path.lower().endswith('.pdf')
    
    def load(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Load PDF document and extract text and metadata"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        try:
            with open(file_path, 'rb') as file:
                return self._process_pdf(file, os.path.basename(file_path), os.path.getsize(file_path))
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise
    
    def load_from_bytes(self, file_bytes: BinaryIO, filename: str) -> Tuple[str, DocumentMetadata]:
        """Load PDF document from bytes"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        try:
            # Get file size by seeking to end and back
            current_pos = file_bytes.tell()
            file_bytes.seek(0, os.SEEK_END)
            file_size = file_bytes.tell()
            file_bytes.seek(current_pos)  # Reset position
            
            return self._process_pdf(file_bytes, filename, file_size)
        except Exception as e:
            logger.error(f"Error loading PDF from bytes: {str(e)}")
            raise
    
    def _process_pdf(self, file_obj: BinaryIO, filename: str, file_size: int) -> Tuple[str, DocumentMetadata]:
        """Process PDF file object"""
        pdf_reader = PyPDF2.PdfReader(file_obj)
        
        # Extract text
        text_content = ""
        pages = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            pages.append(page_text)
            text_content += page_text + "\n\n"
        
        # Extract metadata
        info = pdf_reader.metadata
        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.PDF,
            page_count=len(pdf_reader.pages),
            size_bytes=file_size,
        )
        
        # Add available metadata if present
        if info:
            metadata.author = info.get('/Author', '')
            metadata.creation_date = info.get('/CreationDate', '')
            metadata.modified_date = info.get('/ModDate', '')
            metadata.title = info.get('/Title', '')
        
        return text_content, metadata


class DocxLoader(DocumentLoader):
    """DOCX document loader"""
    
    def can_load(self, file_path: str) -> bool:
        """Check if the file is a DOCX"""
        return file_path.lower().endswith('.docx')
    
    def load(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Load DOCX document and extract text and metadata"""
        if docx is None:
            raise ImportError("python-docx is required for DOCX processing")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract metadata
            metadata = DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type=DocumentType.DOCX,
                page_count=len(doc.sections),
                size_bytes=os.path.getsize(file_path),
            )
            
            # Add core properties if available
            try:
                core_props = doc.core_properties
                metadata.author = core_props.author or ""
                metadata.creation_date = str(core_props.created) if core_props.created else ""
                metadata.modified_date = str(core_props.modified) if core_props.modified else ""
                metadata.title = core_props.title or ""
            except:
                pass
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise
    
    def load_from_bytes(self, file_bytes: BinaryIO, filename: str) -> Tuple[str, DocumentMetadata]:
        """Load DOCX document from bytes"""
        if docx is None:
            raise ImportError("python-docx is required for DOCX processing")
        
        try:
            # Get file size
            current_pos = file_bytes.tell()
            file_bytes.seek(0, os.SEEK_END)
            file_size = file_bytes.tell()
            file_bytes.seek(current_pos)  # Reset position
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(file_bytes.read())
                temp_path = temp_file.name
            
            # Process the temporary file
            try:
                result = self.load(temp_path)
                # Update filename in metadata
                result[1].filename = filename
                result[1].size_bytes = file_size
                return result
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
        except Exception as e:
            logger.error(f"Error loading DOCX from bytes: {str(e)}")
            raise


class TextLoader(DocumentLoader):
    """Plain text document loader"""
    
    def can_load(self, file_path: str) -> bool:
        """Check if the file is a text file"""
        return file_path.lower().endswith(('.txt', '.md', '.csv', '.json', '.xml', '.html'))
    
    def load(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Load text document and extract content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Count pages (approximate by newlines)
            page_count = max(1, text_content.count('\n') // 40)
            
            metadata = DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type=DocumentType.TXT,
                page_count=page_count,
                size_bytes=os.path.getsize(file_path),
            )
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error loading text file: {str(e)}")
            raise
    
    def load_from_bytes(self, file_bytes: BinaryIO, filename: str) -> Tuple[str, DocumentMetadata]:
        """Load text document from bytes"""
        try:
            # Get file size
            current_pos = file_bytes.tell()
            file_bytes.seek(0, os.SEEK_END)
            file_size = file_bytes.tell()
            file_bytes.seek(current_pos)  # Reset position
            
            # Read content
            text_content = file_bytes.read().decode('utf-8')
            
            # Count pages (approximate by newlines)
            page_count = max(1, text_content.count('\n') // 40)
            
            metadata = DocumentMetadata(
                filename=filename,
                file_type=DocumentType.TXT,
                page_count=page_count,
                size_bytes=file_size,
            )
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error loading text from bytes: {str(e)}")
            raise


class ImageLoader(DocumentLoader):
    """Image document loader with OCR capabilities"""
    
    def __init__(self, ocr_engine: str = 'tesseract', languages: List[str] = None):
        """Initialize image loader with OCR engine"""
        self.ocr_engine = ocr_engine.lower()
        self.languages = languages or ['en']
        
        # Initialize OCR engines
        if self.ocr_engine == 'easyocr' and easyocr is not None:
            self.reader = easyocr.Reader(self.languages)
        elif self.ocr_engine == 'tesseract' and pytesseract is not None:
            # Tesseract is initialized on demand
            pass
        else:
            logger.warning(f"OCR engine {ocr_engine} not available. Falling back to basic image metadata.")
    
    def can_load(self, file_path: str) -> bool:
        """Check if the file is an image"""
        return file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'))
    
    def load(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Load image and extract text using OCR"""
        if Image is None:
            raise ImportError("PIL is required for image processing")
        
        try:
            # Load image
            img = Image.open(file_path)
            
            # Extract text using OCR
            text_content = self._perform_ocr(img)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type=DocumentType.IMAGE,
                page_count=1,
                size_bytes=os.path.getsize(file_path),
            )
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            raise
    
    def load_from_bytes(self, file_bytes: BinaryIO, filename: str) -> Tuple[str, DocumentMetadata]:
        """Load image from bytes and extract text using OCR"""
        if Image is None:
            raise ImportError("PIL is required for image processing")
        
        try:
            # Get file size
            current_pos = file_bytes.tell()
            file_bytes.seek(0, os.SEEK_END)
            file_size = file_bytes.tell()
            file_bytes.seek(current_pos)  # Reset position
            
            # Load image
            img = Image.open(io.BytesIO(file_bytes.read()))
            
            # Extract text using OCR
            text_content = self._perform_ocr(img)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=filename,
                file_type=DocumentType.IMAGE,
                page_count=1,
                size_bytes=file_size,
            )
            
            return text_content, metadata
        except Exception as e:
            logger.error(f"Error processing image from bytes: {str(e)}")
            raise
    
    def _perform_ocr(self, image: 'Image') -> str:
        """Perform OCR on the image"""
        if self.ocr_engine == 'easyocr' and easyocr is not None:
            # EasyOCR processing
            results = self.reader.readtext(np.array(image))
            return "\n".join([text for _, text, _ in results])
        
        elif self.ocr_engine == 'tesseract' and pytesseract is not None:
            # Tesseract processing
            lang_str = '+'.join(self.languages)
            return pytesseract.image_to_string(image, lang=lang_str)
        
        else:
            return "[OCR not available - install pytesseract or easyocr]"


class OCREngine:
    """OCR engine for text extraction from images"""
    
    def __init__(self, engine: str = 'tesseract', languages: List[str] = None):
        """Initialize OCR engine"""
        self.engine = engine.lower()
        self.languages = languages or ['en']
        self.image_loader = ImageLoader(ocr_engine=engine, languages=languages)
    
    def extract_text(self, image_path: str) -> str:
        """Extract text from image file"""
        text, _ = self.image_loader.load(image_path)
        return text
    
    def extract_text_from_bytes(self, image_bytes: BinaryIO, filename: str = "image.png") -> str:
        """Extract text from image bytes"""
        text, _ = self.image_loader.load_from_bytes(image_bytes, filename)
        return text
    
    def extract_text_from_pdf(self, pdf_path: str) -> List[str]:
        """Extract text from PDF using OCR"""
        if convert_from_path is None:
            raise ImportError("pdf2image is required for PDF OCR")
        
        try:
            # Convert PDF to images
            images = convert_from_path(pdf_path)
            
            # Extract text from each page
            results = []
            for img in images:
                text = self._perform_ocr(img)
                results.append(text)
            
            return results
        except Exception as e:
            logger.error(f"Error in PDF OCR: {str(e)}")
            raise
    
    def _perform_ocr(self, image: 'Image') -> str:
        """Perform OCR on the image"""
        return self.image_loader._perform_ocr(image)


class NLPProcessor:
    """Natural Language Processing for document text"""
    
    def __init__(self, model: str = 'en_core_web_sm'):
        """Initialize NLP processor with spaCy model"""
        if spacy is None:
            raise ImportError("spaCy is required for NLP processing")
        
        try:
            self.nlp = spacy.load(model)
        except:
            logger.warning(f"spaCy model {model} not found. You may need to download it with: python -m spacy download {model}")
            self.nlp = None
    
    def process_text(self, text: str) -> Dict[str, Any]:
        """Process text with NLP and return extracted information"""
        if self.nlp is None:
            return {"error": "spaCy model not available"}
        
        # Process text with spaCy
        doc = self.nlp(text)
        
        # Extract entities
        entities = {}
        for ent in doc.ents:
            if ent.label_ not in entities:
                entities[ent.label_] = []
            entities[ent.label_].append(ent.text)
        
        # Extract keywords (nouns and proper nouns)
        keywords = [token.text for token in doc if token.pos_ in ('NOUN', 'PROPN') and not token.is_stop]
        
        # Basic sentiment analysis (very simple approach)
        sentiment_words = {
            'positive': ['good', 'great', 'excellent', 'best', 'happy', 'positive', 'success'],
            'negative': ['bad', 'worst', 'poor', 'terrible', 'sad', 'negative', 'fail']
        }
        
        sentiment = {'positive': 0, 'negative': 0, 'neutral': 0}
        
        for token in doc:
            lower_token = token.text.lower()
            if lower_token in sentiment_words['positive']:
                sentiment['positive'] += 1
            elif lower_token in sentiment_words['negative']:
                sentiment['negative'] += 1
            else:
                sentiment['neutral'] += 1
        
        # Normalize sentiment scores
        total = sum(sentiment.values())
        if total > 0:
            for key in sentiment:
                sentiment[key] = sentiment[key] / total
        
        return {
            "entities": entities,
            "keywords": keywords[:20],  # Limit to top 20 keywords
            "sentiment": sentiment
        }


class DocumentClassifier:
    """Document classification using ML techniques"""
    
    def __init__(self, model_type: str = 'tfidf_svm'):
        """Initialize document classifier"""
        if np is None or TfidfVectorizer is None:
            raise ImportError("scikit-learn is required for document classification")
        
        self.model_type = model_type
        self.model = None
        self.classes = []
        
        # Initialize model based on type
        if model_type == 'tfidf_svm':
            self.model = Pipeline([
                ('vectorizer', TfidfVectorizer(max_features=10000)),
                ('classifier', SVC(kernel='linear', probability=True))
            ])
        elif model_type == 'tfidf_nb':
            self.model = Pipeline([
                ('vectorizer', TfidfVectorizer(max_features=10000)),
                ('classifier', MultinomialNB())
            ])
        else:
            raise ValueError(f"Unsupported model type: {model_type}")
    
    def train(self, texts: List[str], labels: List[str]) -> None:
        """Train the classifier with document texts and labels"""
        if self.model is None:
            raise ValueError("Model not initialized")
        
        self.classes = list(set(labels))
        self.model.fit(texts, labels)
    
    def classify(self, text: str) -> Dict[str, float]:
        """Classify a document and return class probabilities"""
        if self.model is None or not self.classes:
            return {"error": "Model not trained"}
        
        try:
            # Get probabilities for each class
            probs = self.model.predict_proba([text])[0]
            
            # Create dictionary of class -> probability
            result = {cls: float(prob) for cls, prob in zip(self.classes, probs)}
            
            return result
        except Exception as e:
            logger.error(f"Classification error: {str(e)}")
            return {"error": str(e)}
    
    def unsupervised_classify(self, texts: List[str], n_clusters: int = 5) -> List[int]:
        """Perform unsupervised classification (clustering)"""
        if TfidfVectorizer is None or KMeans is None:
            raise ImportError("scikit-learn is required for clustering")
        
        # Create TF-IDF features
        vectorizer = TfidfVectorizer(max_features=10000)
        X = vectorizer.fit_transform(texts)
        
        # Perform K-means clustering
        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
        clusters = kmeans.fit_predict(X)
        
        return clusters.tolist()


class InformationExtractor:
    """Extract structured information from documents"""
    
    def __init__(self, use_ai: bool = False):
        """Initialize information extractor"""
        self.use_ai = use_ai
        self.nlp_processor = None
        self.qa_pipeline = None
        
        # Initialize NLP processor if spaCy is available
        if spacy is not None:
            try:
                self.nlp_processor = NLPProcessor()
            except:
                logger.warning("Failed to initialize NLP processor")
        
        # Initialize QA pipeline if transformers is available and AI is enabled
        if use_ai and pipeline is not None:
            try:
                self.qa_pipeline = pipeline("question-answering")
            except:
                logger.warning("Failed to initialize QA pipeline")
    
    def extract_key_value_pairs(self, text: str) -> Dict[str, str]:
        """Extract key-value pairs from text"""
        # Simple regex-based extraction
        kv_pattern = r'([A-Za-z0-9_\- ]+)[\s]*:[\s]*([^\n]+)'
        matches = re.findall(kv_pattern, text)
        
        result = {}
        for key, value in matches:
            key = key.strip()
            value = value.strip()
            if key and value:
                result[key] = value
        
        return result
    
    def extract_tables(self, text: str) -> List[Dict[str, Any]]:
        """Extract tables from text (simplified)"""
        # This is a simplified implementation
        # Real table extraction would require more sophisticated techniques
        
        tables = []
        
        # Look for potential table sections
        table_sections = re.split(r'\n\s*\n', text)
        
        for section in table_sections:
            lines = section.strip().split('\n')
            if len(lines) >= 3:  # At least header, separator, and one data row
                # Check if this looks like a table (has consistent delimiters)
                if all('|' in line for line in lines[:3]) or all('\t' in line for line in lines[:3]):
                    delimiter = '|' if '|' in lines[0] else '\t'
                    
                    # Process header
                    header = [cell.strip() for cell in lines[0].split(delimiter) if cell.strip()]
                    
                    # Process rows
                    rows = []
                    for line in lines[1:]:
                        if delimiter in line:
                            cells = [cell.strip() for cell in line.split(delimiter) if cell.strip()]
                            if len(cells) == len(header):
                                rows.append(dict(zip(header, cells)))
                    
                    if rows:
                        tables.append({
                            "header": header,
                            "rows": rows
                        })
        
        return tables
    
    def extract_form_fields(self, text: str) -> Dict[str, str]:
        """Extract form fields from text"""
        # Similar to key-value extraction but looking for form-like patterns
        form_patterns = [
            r'([A-Za-z0-9_\- ]+)[\s]*:[\s]*([^\n]+)',  # Key: Value
            r'([A-Za-z0-9_\- ]+)[\s]*=[\s]*([^\n]+)',  # Key = Value
            r'([A-Za-z0-9_\- ]+)[\s]*\[([^\]]+)\]',    # Key [Value]
            r'([A-Za-z0-9_\- ]+)[\s]*\(([^\)]+)\)'     # Key (Value)
        ]
        
        result = {}
        for pattern in form_patterns:
            matches = re.findall(pattern, text)
            for key, value in matches:
                key = key.strip()
                value = value.strip()
                if key and value and key not in result:
                    result[key] = value
        
        return result
    
    def answer_questions(self, text: str, questions: List[str]) -> Dict[str, str]:
        """Answer questions about the document using AI"""
        if not self.use_ai or self.qa_pipeline is None:
            return {"error": "QA pipeline not available"}
        
        answers = {}
        for question in questions:
            try:
                result = self.qa_pipeline(question=question, context=text)
                answers[question] = result["answer"]
            except Exception as e:
                logger.error(f"Error answering question: {str(e)}")
                answers[question] = f"Error: {str(e)}"
        
        return answers


class AIEnhancer:
    """AI-based document enhancement and analysis"""
    
    def __init__(self):
        """Initialize AI enhancer"""
        self.summarizer = None
        self.qa_pipeline = None
        
        # Initialize AI pipelines if transformers is available
        if pipeline is not None:
            try:
                self.summarizer = pipeline("summarization")
            except:
                logger.warning("Failed to initialize summarization pipeline")
            
            try:
                self.qa_pipeline = pipeline("question-answering")
            except:
                logger.warning("Failed to initialize QA pipeline")
    
    def summarize(self, text: str, max_length: int = 150) -> str:
        """Generate a summary of the document"""
        if self.summarizer is None:
            return "Summarization not available"
        
        try:
            # Split text into chunks if it's too long
            max_chunk_length = 1000
            if len(text) > max_chunk_length:
                chunks = [text[i:i+max_chunk_length] for i in range(0, len(text), max_chunk_length)]
                summaries = []
                
                for chunk in chunks:
                    summary = self.summarizer(chunk, max_length=max_length // len(chunks), 
                                             min_length=20, do_sample=False)
                    summaries.append(summary[0]['summary_text'])
                
                return " ".join(summaries)
            else:
                summary = self.summarizer(text, max_length=max_length, min_length=30, do_sample=False)
                return summary[0]['summary_text']
        except Exception as e:
            logger.error(f"Summarization error: {str(e)}")
            return f"Summarization error: {str(e)}"
    
    def answer_question(self, text: str, question: str) -> str:
        """Answer a specific question about the document"""
        if self.qa_pipeline is None:
            return "QA not available"
        
        try:
            # Split text into chunks if it's too long and find the most relevant chunk
            max_chunk_length = 1000
            if len(text) > max_chunk_length:
                chunks = [text[i:i+max_chunk_length] for i in range(0, len(text), max_chunk_length)]
                
                # Simple relevance scoring - count question words in each chunk
                question_words = set(question.lower().split())
                scores = []
                
                for chunk in chunks:
                    chunk_lower = chunk.lower()
                    score = sum(1 for word in question_words if word in chunk_lower)
                    scores.append(score)
                
                # Use the most relevant chunk
                best_chunk = chunks[scores.index(max(scores))]
                result = self.qa_pipeline(question=question, context=best_chunk)
            else:
                result = self.qa_pipeline(question=question, context=text)
            
            return result["answer"]
        except Exception as e:
            logger.error(f"QA error: {str(e)}")
            return f"QA error: {str(e)}"
    
    def generate_keywords(self, text: str, count: int = 10) -> List[str]:
        """Generate keywords for the document"""
        if spacy is None:
            return ["Keywords not available"]
        
        try:
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text)
            
            # Extract nouns and proper nouns
            keywords = {}
            for token in doc:
                if token.pos_ in ('NOUN', 'PROPN') and not token.is_stop and len(token.text) > 2:
                    keywords[token.text.lower()] = keywords.get(token.text.lower(), 0) + 1
            
            # Sort by frequency
            sorted_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
            
            return [kw for kw, _ in sorted_keywords[:count]]
        except Exception as e:
            logger.error(f"Keyword generation error: {str(e)}")
            return [f"Keyword error: {str(e)}"]


class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self, ocr_engine: str = 'tesseract', languages: List[str] = None, 
                 use_ai: bool = False, model_path: str = None):
        """Initialize document processor"""
        self.languages = languages or ['en']
        self.use_ai = use_ai
        self.model_path = model_path
        
        # Initialize document loaders
        self.loaders = [
            PDFLoader(),
            DocxLoader(),
            TextLoader(),
            ImageLoader(ocr_engine=ocr_engine, languages=self.languages)
        ]
        
        # Initialize processing components
        self.ocr_engine = OCREngine(engine=ocr_engine, languages=self.languages)
        
        try:
            self.nlp_processor = NLPProcessor() if spacy is not None else None
        except:
            logger.warning("Failed to initialize NLP processor")
            self.nlp_processor = None
        
        try:
            self.classifier = DocumentClassifier() if np is not None else None
        except:
            logger.warning("Failed to initialize document classifier")
            self.classifier = None
        
        self.info_extractor = InformationExtractor(use_ai=use_ai)
        
        try:
            self.ai_enhancer = AIEnhancer() if use_ai else None
        except:
            logger.warning("Failed to initialize AI enhancer")
            self.ai_enhancer = None
    
    def get_loader_for_file(self, file_path: str) -> DocumentLoader:
        """Get appropriate loader for file type"""
        for loader in self.loaders:
            if loader.can_load(file_path):
                return loader
        
        raise ValueError(f"Unsupported file type: {file_path}")
    
    def process_file(self, file_path: str, extract_tables: bool = True, 
                    extract_forms: bool = True, classify: bool = True,
                    summarize: bool = True) -> ProcessingResult:
        """Process document file and extract information"""
        try:
            # Get appropriate loader
            loader = self.get_loader_for_file(file_path)
            
            # Load document
            text_content, metadata = loader.load(file_path)
            
            # Create base result
            result = ProcessingResult(
                metadata=metadata,
                text_content=text_content
            )
            
            # Process with NLP if available
            if self.nlp_processor is not None:
                nlp_results = self.nlp_processor.process_text(text_content)
                result.entities = nlp_results.get("entities")
                result.keywords = nlp_results.get("keywords")
                result.sentiment = nlp_results.get("sentiment")
            
            # Extract tables if requested
            if extract_tables:
                result.tables = self.info_extractor.extract_tables(text_content)
            
            # Extract form fields if requested
            if extract_forms:
                result.form_fields = self.info_extractor.extract_form_fields(text_content)
            
            # Classify document if requested
            if classify and self.classifier is not None and hasattr(self.classifier, 'classify'):
                # This would normally use a pre-trained model
                # Here we're just using a placeholder
                result.classification = {"placeholder": 1.0}
            
            # Generate summary if requested
            if summarize and self.ai_enhancer is not None:
                result.summary = self.ai_enhancer.summarize(text_content)
            
            return result
        
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return ProcessingResult(
                metadata=DocumentMetadata(
                    filename=os.path.basename(file_path),
                    file_type=DocumentType.UNKNOWN,
                    size_bytes=os.path.getsize(file_path) if os.path.exists(file_path) else 0
                ),
                text_content="",
                error=str(e)
            )
    
    def process_bytes(self, file_bytes: BinaryIO, filename: str, 
                     extract_tables: bool = True, extract_forms: bool = True,
                     classify: bool = True, summarize: bool = True) -> ProcessingResult:
        """Process document from bytes and extract information"""
        # Determine file type from extension
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            # Find appropriate loader
            loader = None
            for l in self.loaders:
                if l.can_load(filename):
                    loader = l
                    break
            
            if loader is None:
                raise ValueError(f"Unsupported file type: {filename}")
            
            # Load document
            text_content, metadata = loader.load_from_bytes(file_bytes, filename)
            
            # Create base result
            result = ProcessingResult(
                metadata=metadata,
                text_content=text_content
            )
            
            # Process with NLP if available
            if self.nlp_processor is not None:
                nlp_results = self.nlp_processor.process_text(text_content)
                result.entities = nlp_results.get("entities")
                result.keywords = nlp_results.get("keywords")
                result.sentiment = nlp_results.get("sentiment")
            
            # Extract tables if requested
            if extract_tables:
                result.tables = self.info_extractor.extract_tables(text_content)
            
            # Extract form fields if requested
            if extract_forms:
                result.form_fields = self.info_extractor.extract_form_fields(text_content)
            
            # Classify document if requested
            if classify and self.classifier is not None and hasattr(self.classifier, 'classify'):
                # This would normally use a pre-trained model
                # Here we're just using a placeholder
                result.classification = {"placeholder": 1.0}
            
            # Generate summary if requested
            if summarize and self.ai_enhancer is not None:
                result.summary = self.ai_enhancer.summarize(text_content)
            
            return result
        
        except Exception as e:
            logger.error(f"Error processing bytes for {filename}: {str(e)}")
            
            # Get file size
            current_pos = file_bytes.tell()
            file_bytes.seek(0, os.SEEK_END)
            file_size = file_bytes.tell()
            file_bytes.seek(current_pos)  # Reset position
            
            return ProcessingResult(
                metadata=DocumentMetadata(
                    filename=filename,
                    file_type=DocumentType.UNKNOWN,
                    size_bytes=file_size
                ),
                text_content="",
                error=str(e)
            )
    
    def batch_process(self, file_paths: List[str]) -> Dict[str, ProcessingResult]:
        """Process multiple documents"""
        results = {}
        for file_path in file_paths:
            results[file_path] = self.process_file(file_path)
        return results
    
    def train_classifier(self, file_paths: List[str], labels: List[str]) -> None:
        """Train document classifier with labeled examples"""
        if self.classifier is None:
            raise ValueError("Classifier not available")
        
        texts = []
        for file_path in file_paths:
            try:
                loader = self.get_loader_for_file(file_path)
                text, _ = loader.load(file_path)
                texts.append(text)
            except Exception as e:
                logger.error(f"Error loading {file_path} for training: {str(e)}")
                texts.append("")  # Add empty text to maintain alignment with labels
        
        self.classifier.train(texts, labels)
    
    def extract_specific_info(self, text: str, info_type: str) -> Dict[str, Any]:
        """Extract specific type of information from text"""
        if info_type == 'entities' and self.nlp_processor is not None:
            nlp_results = self.nlp_processor.process_text(text)
            return {"entities": nlp_results.get("entities", {})}
        
        elif info_type == 'key_value':
            return {"key_value": self.info_extractor.extract_key_value_pairs(text)}
        
        elif info_type == 'tables':
            return {"tables": self.info_extractor.extract_tables(text)}
        
        elif info_type == 'form_fields':
            return {"form_fields": self.info_extractor.extract_form_fields(text)}
        
        elif info_type == 'summary' and self.ai_enhancer is not None:
            return {"summary": self.ai_enhancer.summarize(text)}
        
        else:
            return {"error": f"Unsupported info type: {info_type}"}
    
    def answer_questions(self, text: str, questions: List[str]) -> Dict[str, str]:
        """Answer questions about document content"""
        if self.ai_enhancer is None:
            return {"error": "AI enhancer not available"}
        
        answers = {}
        for question in questions:
            answers[question] = self.ai_enhancer.answer_question(text, question)
        
        return answers


# Example usage
if __name__ == "__main__":
    # Initialize processor
    processor = DocumentProcessor(ocr_engine='tesseract', languages=['en'], use_ai=False)
    
    # Process a file
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        result = processor.process_file(file_path)
        print(json.dumps(result.to_dict(), indent=2, default=str))
    else:
        print("Usage: python document_processing.py <file_path>")

